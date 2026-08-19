"""
Parses Section 8 of each SOP .docx into category/subcategory chunks
and (re)builds the Chroma vector database.
"""

import re
import shutil
from pathlib import Path

from docx import Document as WordDocument

import chromadb
from langchain_core.documents import Document
from langchain_chroma import Chroma

from client import embeddings
from rag.metadata import build_search_fields


# ---------------------------------------------------------
# Configuration
# ---------------------------------------------------------
SOP_DIR = Path(__file__).resolve().parent.parent / "data" / "sop"
VECTOR_DB_DIR = Path(__file__).resolve().parent.parent / "vectordb"

COLLECTION_NAME = "agent_guidelines"
TARGET_SECTION = "8. Parameter-by-Parameter Agent Guidelines"


# ---------------------------------------------------------
# Helper functions
# ---------------------------------------------------------
def get_heading_level(paragraph):
    """Returns Word heading level (1,2,3...) or None for normal paragraphs."""
    style_name = paragraph.style.name.strip().lower()
    match = re.search(r"heading\s*(\d+)", style_name)
    return int(match.group(1)) if match else None


def is_category_header(text: str) -> bool:
    """Detects 'EXPERTISE Category', 'PROFESSIONALISM Category', etc."""
    return bool(re.match(r"^.+?\s+Category$", text.strip(), re.IGNORECASE))


def is_numbered_section_header(text: str) -> bool:
    """Detects top-level numbered sections: '8. ...', '9. ...', '10. ...'."""
    return bool(re.match(r"^\d+\.\s+.+$", text.strip()))


def is_target_section_header(text: str) -> bool:
    """Detects Section 8 header exactly."""
    normalized_text = re.sub(r"\s+", " ", text.strip()).lower()
    return normalized_text == TARGET_SECTION.lower()


# ---------------------------------------------------------
# Parse one DOCX file
# ---------------------------------------------------------
def parse_docx_by_headers(docx_path: Path) -> list[Document]:
    """Extracts Section 8 -> one Document per subcategory."""
    word_doc = WordDocument(str(docx_path))

    chunks: list[Document] = []
    inside_target_section = False
    current_category = None
    current_subcategory = None
    current_content: list[str] = []

    def save_current_subcategory():
        if not (current_category and current_subcategory and current_content):
            return

        content = "\n".join(current_content).strip()
        search_fields = build_search_fields(current_category, current_subcategory)

        chunks.append(
            Document(
                page_content=content,
                metadata={
                    "source": str(docx_path),
                    "section_name": TARGET_SECTION,
                    "category_display": current_category,
                    "subcategory_display": current_subcategory,
                    **search_fields,   # category, subcategory, category_combined
                    "chunk_type": "subcategory",
                },
            )
        )

    for paragraph in word_doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        heading_level = get_heading_level(paragraph)

        # ---- Identify Section 8 start ----
        if is_target_section_header(text):
            inside_target_section = True
            continue

        # ---- Stop at next top-level numbered section ----
        if inside_target_section and is_numbered_section_header(text):
            if not is_target_section_header(text):
                break

        if not inside_target_section:
            continue

        # ---- Category header ----
        if is_category_header(text):
            save_current_subcategory()
            current_category = re.sub(
                r"\s+Category$", "", text, flags=re.IGNORECASE
            ).strip().upper()
            current_subcategory = None
            current_content = []
            continue

        # ---- Subcategory header (Heading 3 exactly) ----
        if current_category and heading_level is not None and heading_level == 3:
            save_current_subcategory()
            current_subcategory = text
            current_content = []
            continue

        # ---- Normal content ----
        if current_category and current_subcategory:
            current_content.append(text)

    save_current_subcategory()
    return chunks


# ---------------------------------------------------------
# Load all DOCX files
# ---------------------------------------------------------
def build_documents() -> list[Document]:
    docx_files = list(SOP_DIR.glob("*.docx"))
    if not docx_files:
        raise FileNotFoundError(f"No DOCX files found in: {SOP_DIR}")

    documents: list[Document] = []
    for docx_file in docx_files:
        file_chunks = parse_docx_by_headers(docx_file)
        documents.extend(file_chunks)
        print(f"{docx_file.name}: {len(file_chunks)} subcategory chunks created")

    if not documents:
        raise ValueError(
            "No Section 8 subcategory chunks were created. "
            "Check the section name and Word heading styles."
        )
    return documents


# ---------------------------------------------------------
# Rebuild Chroma database
# ---------------------------------------------------------
def rebuild_vector_db(documents: list[Document]) -> Chroma:
    if VECTOR_DB_DIR.exists():
        shutil.rmtree(VECTOR_DB_DIR)

    client = chromadb.PersistentClient(path=str(VECTOR_DB_DIR))

    db = Chroma.from_documents(
        documents=documents,
        embedding=embeddings,
        collection_name=COLLECTION_NAME,
        client=client,
    )
    return db


# ---------------------------------------------------------
# Main
# ---------------------------------------------------------
if __name__ == "__main__":
    documents = build_documents()
    rebuild_vector_db(documents)

    print("\nVector DB Created Successfully")
    print(f"Location: {VECTOR_DB_DIR}")
    print(f"Total chunks: {len(documents)}")

    for index, chunk in enumerate(documents, start=1):
        print("\n" + "=" * 80)
        print(f"CHUNK {index}")
        print("Category   :", chunk.metadata["category"])
        print("Subcategory:", chunk.metadata["subcategory"])
        print("Combined   :", chunk.metadata["category_combined"])
        print("Source     :", chunk.metadata["source"])
        print("\nContent:")
        print(chunk.page_content)