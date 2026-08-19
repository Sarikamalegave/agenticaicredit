# utils/savedocs.py
"""
DOCX writers for PTK, Assessment, Feedback, and Escalation reports.
All inputs are plain dicts/lists (workflow calls .model_dump() before saving).
"""
from docx import Document


# ----------------------------------------------------------------------
# Generic recursive renderer
# ----------------------------------------------------------------------
def _render_value(doc, value, level: int = 1):
    level = min(level, 4)  # docx supports heading levels 1-9, keep it tidy
    if isinstance(value, dict):
        for k, v in value.items():
            doc.add_heading(str(k).replace("_", " ").title(), level=level)
            _render_value(doc, v, level + 1)
    elif isinstance(value, list):
        for item in value:
            if isinstance(item, (dict, list)):
                _render_value(doc, item, level + 1)
            else:
                doc.add_paragraph(str(item), style="List Bullet")
    else:
        doc.add_paragraph(str(value) if value is not None else "")


# ----------------------------------------------------------------------
# PTK report
# ----------------------------------------------------------------------
def _write_ptk_docx(ptk_list, path):
    """
    ptk_list: [
      {"ptk": {...PTK fields...}, "parameter": "...", "subparameter": "...", "sop_context": "..."},
      ...
    ]
    """
    doc = Document()
    doc.add_heading("Personalized Training Kit (PTK)", level=0)

    if not ptk_list:
        doc.add_paragraph("No PTK entries generated.")
        doc.save(str(path))
        return

    for idx, entry in enumerate(ptk_list, 1):
        parameter = entry.get("parameter", "")
        subparameter = entry.get("subparameter", "")
        doc.add_heading(f"{idx}. {parameter} - {subparameter}", level=1)

        ptk = entry.get("ptk", {})
        if isinstance(ptk, dict):
            for key, value in ptk.items():
                doc.add_heading(str(key).replace("_", " ").title(), level=2)
                if isinstance(value, list):
                    for item in value:
                        doc.add_paragraph(str(item), style="List Bullet")
                else:
                    doc.add_paragraph(str(value) if value is not None else "")
        else:
            doc.add_paragraph(str(ptk))

    doc.save(str(path))


# ----------------------------------------------------------------------
# Assessment report
# ----------------------------------------------------------------------
def _write_assessment_docx(assessment_list, path):
    """
    assessment_list: [
      {"parameter": "...", "subparameter": "...", "questions": [ {question, options, correct_answer, explanation}, ... ]},
      ...
    ]
    """
    doc = Document()
    doc.add_heading("Knowledge Assessment", level=0)

    if not assessment_list:
        doc.add_paragraph("No assessment generated.")
        doc.save(str(path))
        return

    for idx, assessment in enumerate(assessment_list, 1):
        if not isinstance(assessment, dict):
            doc.add_paragraph(str(assessment))
            continue

        parameter = assessment.get("parameter", "")
        subparameter = assessment.get("subparameter", "")
        title = f"{idx}. {parameter} - {subparameter}".strip(" -")
        doc.add_heading(title or f"Assessment {idx}", level=1)

        questions = assessment.get("questions", [])
        if isinstance(questions, list) and questions:
            for qi, q in enumerate(questions, 1):
                if not isinstance(q, dict):
                    doc.add_paragraph(str(q))
                    continue
                doc.add_heading(f"Q{qi}. {q.get('question', '')}", level=2)

                options = q.get("options", [])
                for opt in options:
                    doc.add_paragraph(str(opt), style="List Bullet")

                if q.get("correct_answer"):
                    doc.add_paragraph(f"Correct Answer: {q.get('correct_answer')}")
                if q.get("explanation"):
                    doc.add_paragraph(f"Explanation: {q.get('explanation')}")
        else:
            # fall back to generic rendering for unexpected shapes
            _render_value(doc, assessment, level=2)

    doc.save(str(path))


# ----------------------------------------------------------------------
# Generic report (Feedback / Escalation)
# ----------------------------------------------------------------------
def _write_report_docx(data, path, title: str):
    """
    Generic writer for Feedback / Escalation dicts.
    Handles nested lists/dicts (e.g. coaching_reference, fatal_breaches).
    """
    doc = Document()
    doc.add_heading(title, level=0)
    _render_value(doc, data, level=1)
    doc.save(str(path))